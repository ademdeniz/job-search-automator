# Copyright (c) 2026 Adem Garic. All rights reserved.
# Unauthorized use, copying, or distribution is prohibited. See LICENSE.
"""
Resume Tailor + Cover Letter Generator using Claude API.

Given a job posting and the candidate's resume, generates:
  1. A tailored resume — same content, bullets reworded/reordered to mirror
     the job's language and priorities.
  2. A professional cover letter — specific to the role and company, not generic.

Both are saved as .docx files in output/<slug>/.
"""

import json
import os
import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")

MODEL = "claude-sonnet-4-6"

_TAILOR_SYSTEM_TEMPLATE = textwrap.dedent("""
    You are an expert resume writer and career coach specialising in QA / SDET / test automation roles.
    You will receive a candidate's resume, a job posting, and optionally a company context snippet.

    Your task is to produce TWO documents, returned as a single JSON object (no markdown fences):

    {
      "real_company": "<the actual hiring company name — NOT a job board or aggregator>",
      "tailored_resume": "<full resume text, structured as described below>",
      "cover_letter": "<full cover letter text, plain paragraphs separated by blank lines>"
    }

    For "real_company": always identify the actual employer from the job description.
    If the company field says "RemoteHunter", "Jobgether", "Sundayy", "Scoutit", "LinkedIn", or any
    other aggregator/job board, read the description and return the real hiring company instead.

    TAILORED RESUME structure (use exactly these markers):
    Line 1: Candidate full name (no pronouns)
    Line 2: Contact line (city · email · LinkedIn · GitHub)
    Line 3: blank
    <<<SECTION: PROFESSIONAL SUMMARY>>>
    2-3 sentences tailored to THIS specific job — connect the candidate's background directly to the role's priorities.
    <<<SECTION: CORE SKILLS>>>
    Skill groups as "Category: skill1, skill2, skill3" — one per line, reordered so most relevant to this job comes first.
    <<<SECTION: PROFESSIONAL EXPERIENCE>>>
    For each role:
    TITLE | COMPANY | DATE RANGE
    - bullet (start with strong action verb, mirror job posting language)
    - bullet
    (blank line between roles)
    <<<SECTION: EDUCATION>>>
    Degree — Institution | Year
    <<<SECTION: CERTIFICATIONS>>>
    - Cert name — Issuer (Year)
    <<<SECTION: LANGUAGES>>>
    Languages line

    RESUME BULLET RULES — critical:
    - Use the EXACT metrics, numbers, and project names from the candidate's resume. Never round or genericise.
      If the resume says "20+ test cases", write "20+ test cases". If it says "Appium-based mobile automation
      framework from scratch for iOS and Android", reference that specific framework by name.
    - Surface the most impressive concrete outcomes first (built X, reduced Y, saved Z).
    - Every bullet must start with a strong past-tense action verb and include at least one specific detail.
    - Keep ALL real experience — do not invent anything.
    - Reword bullets to echo the job posting's language naturally, but keep the real specifics intact.
    - Prioritise the most relevant bullets for this job.
    - Preserve actual dates, companies, and titles exactly.
    - {github_instruction}

    COVER LETTER rules:
    - Exactly 3 paragraphs plus a sign-off. No more.
    - NO "Dear Hiring Manager" opener on its own line — skip it entirely.
    - Paragraph 1 (2-3 sentences MAX): Name the company and the role. State one specific reason this
      role fits — tied to their product, mission, or tech stack (use Company Context if available).
      End with one sharp sentence on the candidate's single most relevant qualification.
    - Paragraph 2 (3-4 sentences): One concrete story — the candidate's biggest relevant achievement
      with exact metrics and project names from the resume. Impact first, then what they did. Do not
      list multiple achievements; go deep on one.
    - Paragraph 3 (2-3 sentences): Two or three specific skills or qualifications that map directly
      to the job's stated requirements. Name the technology or method. End with a confident call to
      action referencing the role by name.
    - Sign-off: "Best regards,\\n\\n{name}\\n{email}{linkedin_line}"

    {writing_sample_block}
    HUMAN-SOUNDING TONE — this is critical:
    - Use contractions naturally: "I've", "I'm", "it's", "that's", "I'd", "there's".
    - Vary sentence length intentionally. Mix short punchy sentences with longer ones.
      A short sentence after a long one creates rhythm. Do not make every sentence the same length.
    - Never start two consecutive sentences with "I".
    - Do not use parallel list structures ("X, Y, and Z") — they are an instant AI tell.
      Instead, describe one thing specifically or connect ideas with "and" naturally.
    - Ban these words and phrases entirely: "high-impact", "results-first", "passionate",
      "excited to apply", "I believe", "I feel", "dedicated", "leverage", "utilize",
      "align", "directly reducing", "increase confidence", "built for", "designed to".
    - Do not write in perfectly balanced sentences. Real humans leave ideas slightly open.
    - Write as if the candidate is speaking to someone they respect, not performing for a committee.
    - NEVER use em dashes (—). Use a comma, period, or rewrite the sentence instead.

    GLOBAL FORMATTING RULE: Never use em dashes (—) anywhere in either document.
    They are an immediate AI tell. Rewrite any sentence that would need one.

    Return ONLY the JSON object. No explanation, no markdown fences.
""").strip()


def _build_system_prompt(profile: dict) -> str:
    name     = profile.get("name", "")
    email    = profile.get("email", "")
    linkedin = profile.get("linkedin", "").strip().rstrip("/")
    github   = profile.get("github", "").strip().rstrip("/")
    writing_sample = profile.get("writing_sample", "").strip()

    contact_parts = [c for c in [email, linkedin] if c]
    linkedin_line = " | " + " | ".join(contact_parts[1:]) if len(contact_parts) > 1 else ""

    github_instruction = (
        f"Always include {github} in the contact line."
        if github else "Include the candidate's GitHub URL in the contact line if available."
    )

    if writing_sample:
        writing_sample_block = (
            "VOICE CALIBRATION — the candidate's actual writing (not a job document):\n"
            "```\n"
            f"{writing_sample[:1500]}\n"
            "```\n"
            "Study this carefully. Mirror the candidate's natural rhythm and sentence patterns:\n"
            "- If their sentences are short and declarative, keep them short and declarative.\n"
            "- Match their level of specificity — they name exact places, brands, details.\n"
            "- They don't over-explain observations. State it and move on.\n"
            "- Their humor is dry and understated. If something is slightly absurd, they just note it.\n"
            "- They mix short punchy statements with longer observational ones.\n"
            "- Do not impose a corporate voice on top of their natural one.\n"
        )
    else:
        writing_sample_block = ""

    return (
        _TAILOR_SYSTEM_TEMPLATE
        .replace("{name}", name)
        .replace("{email}", email)
        .replace("{linkedin_line}", linkedin_line)
        .replace("{github_instruction}", github_instruction)
        .replace("{writing_sample_block}", writing_sample_block)
    )


@dataclass
class TailorResult:
    tailored_resume: str
    cover_letter: str
    output_dir: str
    resume_path: str
    cover_letter_path: str


def _load_resume() -> str:
    """Load resume text from profile.json, falling back to resume.txt."""
    import sys
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
    from storage.profile import load_profile
    profile = load_profile()
    if profile.get("resume", "").strip():
        return profile["resume"].strip()
    # Legacy fallback
    resume_path = os.path.join(os.path.dirname(__file__), "..", "resume.txt")
    if os.path.exists(resume_path):
        with open(resume_path, encoding="utf-8") as f:
            return f.read().strip()
    raise FileNotFoundError("No resume found. Add your resume in the Profile page.")


# Job aggregators — when the stored company is one of these,
# Claude must extract the real employer from the description instead.
AGGREGATORS = {
    "remotehunter", "jobgether", "indeed", "linkedin", "glassdoor",
    "ziprecruiter", "simplyhired", "careerbuilder", "monster", "dice",
    "jobboard", "jobsite", "scoutit", "talent.com", "jobs",
}


def _is_aggregator(company: str) -> bool:
    return company.lower().strip() in AGGREGATORS


def _slug(title: str, company: str) -> str:
    raw = f"{title}_{company}"
    return re.sub(r"[^a-z0-9]+", "_", raw.lower()).strip("_")[:60]


# Domains that host job postings but aren't the actual company site.
_JOB_BOARD_DOMAINS = {
    "greenhouse.io", "boards.greenhouse.io", "lever.co", "jobs.lever.co",
    "linkedin.com", "indeed.com", "dice.com", "remoteok.com",
    "weworkremotely.com", "remoteok.io", "glassdoor.com", "ziprecruiter.com",
}


def _fetch_company_context(job_url: str) -> str:
    """
    Try to fetch a brief description of the company from their own website.
    Returns a clean text snippet (≤1 000 chars) or empty string on failure.
    """
    import requests as _req

    if not job_url:
        return ""

    # Extract domain from job URL
    m = re.search(r"https?://(?:www\.)?([^/]+)", job_url)
    if not m:
        return ""
    domain = m.group(1).lower()

    # Skip job board domains — they don't tell us about the company
    if any(board in domain for board in _JOB_BOARD_DOMAINS):
        return ""

    headers = {"User-Agent": "Mozilla/5.0 (compatible; job-search-bot/1.0)"}

    for url in [f"https://{domain}/about", f"https://www.{domain}/about", f"https://{domain}"]:
        try:
            resp = _req.get(url, timeout=6, headers=headers, allow_redirects=True)
            if resp.status_code != 200:
                continue
            # Strip tags and collapse whitespace
            text = re.sub(r"<[^>]+>", " ", resp.text)
            text = re.sub(r"\s+", " ", text).strip()
            if len(text) > 150:
                return text[:1000]
        except Exception:
            continue

    return ""


def tailor_job(job: dict, resume_text: Optional[str] = None) -> TailorResult:
    if resume_text is None:
        resume_text = _load_resume()

    title       = job.get("title", "")
    company     = job.get("company", "")
    description = job.get("description", "") or ""
    job_url     = job.get("url", "") or ""

    if not description.strip():
        raise ValueError(f"Job {job.get('id')} has no description — cannot tailor. Fetch the description first.")

    # If job was scraped via an aggregator, tell Claude to find the real company
    if _is_aggregator(company):
        company_line = (
            f"Listed via: {company} (aggregator — identify the REAL hiring company "
            f"from the job description and use that company name throughout all output)"
        )
    else:
        company_line = f"Company: {company}"

    # Fetch company context from their own website (best-effort, silent on failure)
    print(f"[Tailor] Fetching company context from {job_url or 'N/A'}…")
    company_context = _fetch_company_context(job_url)
    if company_context:
        print(f"[Tailor] Got {len(company_context)} chars of company context.")
    else:
        print("[Tailor] No company context found — proceeding with job description only.")

    company_context_block = (
        f"\n## Company Context (from their website)\n{company_context}\n"
        if company_context else ""
    )

    # Build system prompt from profile
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
    from storage.profile import load_profile as _load_profile
    profile = _load_profile()
    system_prompt = _build_system_prompt(profile)

    client = anthropic.Anthropic()

    user_msg = textwrap.dedent(f"""
        ## Candidate Resume
        {resume_text}

        ## Job Posting
        Title:   {title}
        {company_line}

        {description[:4000]}
        {company_context_block}
    """).strip()

    print(f"[Tailor] Calling Claude {MODEL} for: {title} @ {company}…")
    message = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_msg}],
    )

    raw = message.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            data = json.loads(m.group())
        else:
            raise ValueError(f"Could not parse JSON from tailor response:\n{raw[:500]}")

    tailored_resume = data.get("tailored_resume", "").strip()
    cover_letter    = data.get("cover_letter", "").strip()
    # Use Claude's identified real company — never an aggregator name
    real_company    = data.get("real_company", "").strip() or company

    # Strip em dashes — an immediate AI tell. Replace with a plain hyphen-minus.
    tailored_resume = tailored_resume.replace("\u2014", "-").replace("\u2013", "-")
    cover_letter    = cover_letter.replace("\u2014", "-").replace("\u2013", "-")

    slug    = _slug(title, real_company)
    out_dir = os.path.join(OUTPUT_DIR, slug)
    os.makedirs(out_dir, exist_ok=True)

    resume_path = os.path.join(out_dir, "resume.docx")
    cl_path     = os.path.join(out_dir, "cover_letter.docx")

    _write_resume_docx(tailored_resume, resume_path)
    _write_cover_letter_docx(cover_letter, title, real_company, cl_path)
    _ats_validate(resume_path)

    if real_company != company:
        print(f"[Tailor] Real company identified: {real_company} (was: {company})")
    print(f"[Tailor] Saved to {out_dir}/")
    return TailorResult(
        tailored_resume=tailored_resume,
        cover_letter=cover_letter,
        output_dir=out_dir,
        resume_path=resume_path,
        cover_letter_path=cl_path,
    )


# ── DOCX helpers ──────────────────────────────────────────────────────────────

DARK      = RGBColor(0x1e, 0x29, 0x3b)   # slate-800 — body text
ACCENT    = RGBColor(0x1d, 0x4e, 0xd8)   # blue-700  — section headers
MUTED     = RGBColor(0x64, 0x74, 0x8b)   # slate-500 — contact / dates
FONT_NAME = "Calibri"


def _font(run, size: float, bold=False, italic=False, color: RGBColor = None):
    run.font.name   = FONT_NAME
    run.font.size   = Pt(size)
    run.bold        = bold
    run.italic      = italic
    if color:
        run.font.color.rgb = color


def _para_fmt(p, space_before=0, space_after=0, line_spacing=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after  = Pt(space_after)
    if line_spacing:
        fmt.line_spacing = Pt(line_spacing)


def _add_hrule(doc):
    """Add a thin horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    _para_fmt(p, space_before=2, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1d4ed8")   # same blue as ACCENT
    pBdr.append(bottom)
    pPr.append(pBdr)


def _ats_validate(path: str):
    """
    Post-process the resume .docx for ATS compatibility:
    - Warn on any tables (ATS parsers often skip table content)
    - Ensure all paragraphs are left-aligned (centred name is fine, rest must be left)
    - Ensure no text boxes (not generated by our writer, but defensive check)
    Prints a summary — does not modify the file if it already passes.
    """
    doc = Document(path)
    issues = []

    if doc.tables:
        issues.append(f"{len(doc.tables)} table(s) found — ATS may skip table content")

    # Check for text boxes in the XML
    from docx.oxml.ns import qn as _qn
    body_xml = doc.element.body.xml
    if "w:txbx" in body_xml:
        issues.append("Text box(es) detected — ATS parsers cannot read text boxes")

    # Check paragraph alignment — only the name (first non-empty para) may be centred
    first_content = True
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if first_content:
            first_content = False
            continue  # name line — any alignment ok
        if para.alignment and para.alignment.name == "CENTER":
            issues.append(f"Centred paragraph found: \"{para.text[:40]}\" — left-align for ATS")
            break  # one warning is enough

    if issues:
        print("[ATS] ⚠️  Potential ATS issues detected:")
        for issue in issues:
            print(f"       • {issue}")
    else:
        print("[ATS] ✅ Resume passes basic ATS formatting checks.")


def _write_resume_docx(text: str, path: str):
    doc = Document()

    # ── Margins ───────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    # Remove default style spacing
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    lines = text.splitlines()
    i = 0
    name_written = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Section marker ────────────────────────────────────────────────────
        m = re.match(r"<<<SECTION:\s*(.+?)>>>", stripped)
        if m:
            _add_hrule(doc)
            p = doc.add_paragraph()
            _para_fmt(p, space_before=4, space_after=2)
            run = p.add_run(m.group(1).upper())
            _font(run, size=9.5, bold=True, color=ACCENT)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not stripped:
            if not name_written:
                i += 1
                continue
            doc.add_paragraph()
            i += 1
            continue

        # ── Name (first non-empty line) ───────────────────────────────────────
        if not name_written:
            p = doc.add_paragraph()
            _para_fmt(p, space_after=2)
            run = p.add_run(stripped)
            _font(run, size=20, bold=True, color=DARK)
            name_written = True
            i += 1
            continue

        # ── Contact line (contains · or @ or linkedin or github) ─────────────
        contact_signals = ("·", "@", "linkedin", "github", "linkedin.com", "github.com")
        if any(s in stripped.lower() for s in contact_signals) and i < 4:
            p = doc.add_paragraph()
            _para_fmt(p, space_after=4)
            run = p.add_run(stripped)
            _font(run, size=9, color=MUTED)
            i += 1
            continue

        # ── Experience role header: TITLE | COMPANY | DATE ───────────────────
        if "|" in stripped and stripped.count("|") >= 1:
            parts = [p.strip() for p in stripped.split("|")]
            p = doc.add_paragraph()
            _para_fmt(p, space_before=6, space_after=1)
            # Title bold, company and date muted
            r_title = p.add_run(parts[0])
            _font(r_title, size=10, bold=True, color=DARK)
            if len(parts) >= 2:
                r_sep = p.add_run("  |  ")
                _font(r_sep, size=10, color=MUTED)
                r_co = p.add_run(parts[1])
                _font(r_co, size=10, color=MUTED)
            if len(parts) >= 3:
                r_sep2 = p.add_run("  |  ")
                _font(r_sep2, size=10, color=MUTED)
                r_date = p.add_run(parts[2])
                _font(r_date, size=10, italic=True, color=MUTED)
            i += 1
            continue

        # ── Bullet ────────────────────────────────────────────────────────────
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _para_fmt(p, space_before=1, space_after=1)
            run = p.add_run(stripped[2:])
            _font(run, size=9.5, color=DARK)
            i += 1
            continue

        # ── Everything else (education line, cert, language) ──────────────────
        p = doc.add_paragraph()
        _para_fmt(p, space_before=1, space_after=1)
        run = p.add_run(stripped)
        _font(run, size=9.5, color=DARK)
        i += 1

    doc.save(path)


def _write_cover_letter_docx(text: str, title: str, company: str, path: str):
    import sys as _sys
    _sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
    from storage.profile import load_profile as _load_profile
    _profile = _load_profile()
    _name    = _profile.get("name", "")
    _email   = _profile.get("email", "")
    _linkedin = _profile.get("linkedin", "").strip().rstrip("/")
    _github   = _profile.get("github", "").strip().rstrip("/")
    _contact_parts = [p for p in [_email, _linkedin, _github] if p]
    _contact_line  = "  ·  ".join(_contact_parts)

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    # ── Letterhead ────────────────────────────────────────────────────────────
    p_name = doc.add_paragraph()
    r = p_name.add_run(_name)
    _font(r, size=16, bold=True, color=DARK)

    p_contact = doc.add_paragraph()
    r2 = p_contact.add_run(_contact_line)
    _font(r2, size=9, color=MUTED)
    _para_fmt(p_contact, space_after=4)

    _add_hrule(doc)

    # Date
    p_date = doc.add_paragraph()
    r_date = p_date.add_run(datetime.now().strftime("%B %d, %Y"))
    _font(r_date, size=10, color=MUTED)
    _para_fmt(p_date, space_before=8, space_after=2)

    # Role line
    p_role = doc.add_paragraph()
    r_role = p_role.add_run(f"Re: {title}  —  {company}")
    _font(r_role, size=10, bold=True, color=ACCENT)
    _para_fmt(p_role, space_after=12)

    # ── Body paragraphs ───────────────────────────────────────────────────────
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        # Sign-off block — preserve line breaks
        if "Best regards" in para or "Adem Garic" in para:
            for line in para.splitlines():
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                bold = "Best regards" in line or line.strip() == "Adem Garic"
                _font(run, size=11, bold=bold, color=DARK)
                _para_fmt(p, space_before=0, space_after=1)
            continue

        p = doc.add_paragraph()
        run = p.add_run(para)
        _font(run, size=11, color=DARK)
        _para_fmt(p, space_before=0, space_after=10, line_spacing=14)

    doc.save(path)
