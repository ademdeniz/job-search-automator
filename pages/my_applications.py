# Copyright (c) 2026 Adem Garic. All rights reserved.
# Unauthorized use, copying, or distribution is prohibited. See LICENSE.
import io
from datetime import datetime, timedelta
from urllib.parse import urlencode

import streamlit as st
from docx import Document as _Document
from docx.shared import Pt as _Pt

from storage.database import (
    get_all_jobs, update_status, delete_job, update_notes, update_interview_prep, stats,
)
from storage.profile import load_profile
from pages.utils import (
    SCORE_COLOR, PIPELINE, PIPELINE_COLOR, match_level, score_badge, fmt_date, claude_call,
    docx_to_pdf, has_libreoffice, render_company_signals,
)

NEXT_STAGE = {
    "applied":      "interviewing",
    "interviewing": "offer",
}
NEXT_LABEL = {
    "applied":      "→ Interviewing",
    "interviewing": "→ Offer",
}


def _render_gmail_scanner(app_jobs: list):
    """Inbox scanner expander — detects responses from applied companies."""
    import gmail_scanner as _gs

    with st.expander("📬 Scan Inbox for Responses", expanded=False):
        if not _gs.is_configured():
            st.markdown(
                "**Gmail setup required.** Connect your Google account to auto-detect "
                "interview invites, rejections, and replies.\n\n"
                "**Steps:**\n"
                "1. Go to [console.cloud.google.com](https://console.cloud.google.com) → "
                "APIs & Services → Library → enable **Gmail API**\n"
                "2. Go to APIs & Services → Credentials → **Create credentials** → "
                "OAuth 2.0 Client ID → Desktop app → Download JSON\n"
                "3. Save the downloaded file as **`credentials.json`** in the app root folder\n"
                "4. Restart the app — a browser tab will open for Google sign-in on first scan"
            )
            return

        auth_note = "" if _gs.is_authenticated() else " *(browser sign-in required on first run)*"
        st.markdown(
            f"Scans your inbox for emails from companies you applied to and detects "
            f"interview invites, rejections, and offers using Claude AI.{auth_note}"
        )

        days_back = st.slider("Look back", 7, 90, 30, step=7,
                              format="%d days", key="gmail_days_back")

        if st.button("📬 Scan Now", type="primary", key="gmail_scan_btn"):
            with st.spinner("Scanning inbox… (opening browser for sign-in if first run)"):
                result = _gs.scan_for_responses(app_jobs, days_back=days_back)

            if result["error"]:
                st.error(result["error"])
                return

            scanned = result["scanned"]
            hits    = result["results"]

            if not hits:
                st.info(
                    f"Scanned {scanned} email(s) — no actionable responses detected. "
                    "Check back after more time has passed, or expand the look-back window."
                )
                return

            st.success(f"Found **{len(hits)} response(s)** in {scanned} email(s) scanned.")
            st.session_state["gmail_scan_results"] = hits

        # ── display results ───────────────────────────────────────────────────
        if "gmail_scan_results" in st.session_state:
            hits = st.session_state["gmail_scan_results"]
            _TYPE_ICON = {
                "interview_invite": "🎯",
                "rejection":        "❌",
                "offer":            "🏆",
                "info_request":     "❓",
                "confirmation":     "✅",
                "other":            "📧",
            }
            _TYPE_COLOR = {
                "interview_invite": "#22c55e",
                "rejection":        "#ef4444",
                "offer":            "#f59e0b",
                "info_request":     "#60a5fa",
                "confirmation":     "#94a3b8",
                "other":            "#64748b",
            }

            for hit in hits:
                icon  = _TYPE_ICON.get(hit["type"], "📧")
                color = _TYPE_COLOR.get(hit["type"], "#94a3b8")
                conf  = int(hit.get("confidence", 0) * 100)

                st.markdown(
                    f'<div style="border:1px solid {color}44;border-radius:8px;'
                    f'padding:12px 16px;margin-bottom:10px;background:{color}0d;">'
                    f'<div style="display:flex;justify-content:space-between;align-items:center;">'
                    f'<span style="color:{color};font-weight:700;font-size:0.95rem;">'
                    f'{icon} {hit["type"].replace("_", " ").title()}'
                    f'</span>'
                    f'<span style="color:#64748b;font-size:0.75rem;">{conf}% confidence</span>'
                    f'</div>'
                    f'<div style="color:#f1f5f9;font-size:0.88rem;margin-top:4px;">'
                    f'<strong>{hit["company"]}</strong> · {hit["subject"]}'
                    f'</div>'
                    f'<div style="color:#94a3b8;font-size:0.8rem;margin-top:2px;">{hit["summary"]}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                suggested = hit.get("suggested_status")
                if suggested:
                    col1, col2 = st.columns([3, 1])
                    with col2:
                        btn_label = f"→ Mark as {suggested.title()}"
                        if st.button(btn_label, key=f"gmail_apply_{hit['job_id']}_{hit['msg_id']}"):
                            from storage.database import update_status
                            update_status(hit["job_id"], suggested)
                            st.toast(f"Status updated to {suggested}!", icon="✅")
                            del st.session_state["gmail_scan_results"]
                            st.rerun()

            if st.button("Clear results", key="gmail_clear"):
                del st.session_state["gmail_scan_results"]
                st.rerun()


def render():
    st.title("📁 My Applications")
    st.caption("Your application history — preserved across fresh searches.")

    # ── filters ───────────────────────────────────────────────────────────────
    f1, f2, f3 = st.columns([2, 1.5, 1.5])
    with f1:
        app_keyword = st.text_input("🔍 Search", placeholder="title, company…", key="app_kw")
    with f2:
        app_status = st.selectbox("Stage", ["All"] + PIPELINE, key="app_status")
    with f3:
        app_sort = st.selectbox(
            "Sort by",
            ["Date applied ↓", "Date applied ↑", "Score ↓", "Score ↑"],
            key="app_sort",
        )

    app_jobs = get_all_jobs(
        status=None if app_status == "All" else app_status,
        keyword=app_keyword or None,
    )
    app_jobs = [j for j in app_jobs if j["status"] in PIPELINE]

    _reverse = "↓" in app_sort
    if app_sort.startswith("Score"):
        app_jobs.sort(key=lambda j: (j.get("score") or -1), reverse=_reverse)
    else:
        app_jobs.sort(key=lambda j: (j.get("applied_at") or ""), reverse=_reverse)

    # ── pipeline summary bar ──────────────────────────────────────────────────
    db_stats = stats()
    _ap = db_stats["by_status"]
    m1, m2, m3 = st.columns(3)
    m1.metric("Applied",      _ap.get("applied", 0))
    m2.metric("Interviewing", _ap.get("interviewing", 0))
    m3.metric("Offer",        _ap.get("offer", 0))
    st.divider()

    if not app_jobs:
        st.info("No applications yet. Mark jobs as 'applied' from the Job Board to track them here.")
        st.stop()

    # ── gmail inbox scanner ───────────────────────────────────────────────────
    _render_gmail_scanner(app_jobs)

    st.caption(f"{len(app_jobs)} application(s)")

    # ── application cards ─────────────────────────────────────────────────────
    for job in app_jobs:
        score    = job.get("score")
        level    = match_level(score)
        s_color  = SCORE_COLOR.get(level, "#94a3b8")
        st_color = PIPELINE_COLOR.get(job["status"], "#94a3b8")

        _chips = "".join([
            f'<span style="background:{s_color}22;color:{s_color};padding:2px 10px;border-radius:99px;font-size:0.8rem;">{score_badge(score)}</span>',
            f'<span style="background:{st_color}22;color:{st_color};padding:2px 10px;border-radius:99px;font-size:0.8rem;font-weight:600;">{job["status"]}</span>',
            f'<span style="color:#64748b;font-size:0.75rem;">&#10003; Applied: {fmt_date(job.get("applied_at", ""))}</span>',
            f'<span style="color:#64748b;font-size:0.75rem;">&#128193; {job["job_type"]}</span>' if job.get("job_type") else "",
            f'<span style="color:#64748b;font-size:0.75rem;">&#128176; {job["salary"]}</span>' if job.get("salary") else "",
            f'<span style="color:#64748b;font-size:0.8rem;">{job["source"]}</span>',
        ])

        with st.container():
            st.markdown(
                f"""<div style="border-left:4px solid {st_color};padding:12px 16px;background:#1e293b;border-radius:6px;margin-bottom:4px;">
<div style="display:flex;justify-content:space-between;align-items:center;">
<div>
<span style="font-size:1.05rem;font-weight:600;color:#f1f5f9;">{job['title']}</span>
&nbsp;<span style="color:#94a3b8;font-size:0.9rem;">{job['company']} · {job['location']}</span>
</div>
<div style="display:flex;gap:8px;align-items:center;flex-wrap:wrap;">{_chips}</div>
</div></div>""",
                unsafe_allow_html=True,
            )

            # ── inline action buttons ─────────────────────────────────────────
            btn_cols = st.columns([1, 1, 1, 4])
            next_stage = NEXT_STAGE.get(job["status"])
            if next_stage:
                if btn_cols[0].button(
                    NEXT_LABEL[job["status"]],
                    key=f"advance_{job['id']}",
                    use_container_width=True,
                ):
                    update_status(job["id"], next_stage)
                    st.rerun()
            if btn_cols[1].button(
                "🗑 Remove",
                key=f"delete_{job['id']}",
                use_container_width=True,
            ):
                delete_job(job["id"])
                st.rerun()

            with st.expander(f"Details — ID {job['id']}", expanded=False):
                d1, d2 = st.columns([2, 1])
                with d1:
                    if job.get("score_reason"):
                        st.markdown(f"**AI Analysis:** {job['score_reason']}")
                    desc = job.get("description") or "No description saved."
                    st.markdown(f"**Description:**\n\n{desc[:1000]}{'…' if len(desc) > 1000 else ''}")
                with d2:
                    _sal = job.get("salary") or ""
                    _est = job.get("salary_estimate") or ""
                    if _sal:
                        st.markdown(f"**Salary:** {_sal}")
                    elif _est:
                        st.markdown(f"**Salary:** ~{_est} *(est.)*")
                    else:
                        st.markdown("**Salary:** N/A")
                    st.markdown(f"**Posted:** {job.get('posted_date') or 'N/A'}")
                    st.markdown(f"**Remote:** {'Yes' if job.get('remote') else 'No'}")
                    if job.get("url"):
                        st.link_button("🔗 Open Job", job["url"])

                    st.divider()
                    render_company_signals(job)

                    st.divider()
                    # Stage corrector — lets you fix mistakes / go back
                    correct_stage = st.selectbox(
                        "Correct stage",
                        PIPELINE,
                        index=PIPELINE.index(job["status"]) if job["status"] in PIPELINE else 0,
                        key=f"correct_{job['id']}",
                    )
                    if correct_stage != job["status"]:
                        if st.button("Save", key=f"correctsave_{job['id']}"):
                            update_status(job["id"], correct_stage)
                            st.rerun()

                    # Notes — loaded from DB, saved on change
                    notes_key = f"notes_{job['id']}"
                    if notes_key not in st.session_state:
                        st.session_state[notes_key] = job.get("notes") or ""
                    new_notes = st.text_area(
                        "Notes (interview prep, contacts, follow-up dates…)",
                        key=notes_key,
                        height=100,
                        placeholder="e.g. Phone screen with Sarah on Apr 10, asked about Appium experience…",
                    )
                    if new_notes != (job.get("notes") or ""):
                        update_notes(job["id"], new_notes)

                    # ── Add to Google Calendar ────────────────────────────────
                    st.markdown("**Schedule Interview**")
                    cal_c1, cal_c2, cal_c3 = st.columns([1.2, 1, 1])
                    with cal_c1:
                        cal_date = st.date_input("Date", key=f"cal_date_{job['id']}", label_visibility="collapsed")
                    with cal_c2:
                        cal_time = st.time_input("Time", key=f"cal_time_{job['id']}", label_visibility="collapsed", step=900)
                    with cal_c3:
                        cal_dur = st.selectbox("Duration", ["30 min", "45 min", "60 min", "90 min"], key=f"cal_dur_{job['id']}", label_visibility="collapsed")

                    _dur_map = {"30 min": 30, "45 min": 45, "60 min": 60, "90 min": 90}
                    _start = datetime.combine(cal_date, cal_time)
                    _end   = _start + timedelta(minutes=_dur_map[cal_dur])
                    _fmt   = "%Y%m%dT%H%M%S"
                    _gc_params = urlencode({
                        "action": "TEMPLATE",
                        "text": f"Interview — {job['title']} @ {job['company']}",
                        "dates": f"{_start.strftime(_fmt)}/{_end.strftime(_fmt)}",
                        "details": f"Job ID: {job['id']}\n{job.get('url', '')}",
                    })
                    _gc_url = f"https://calendar.google.com/calendar/render?{_gc_params}"
                    st.link_button("📅 Add to Google Calendar", _gc_url)

                    # ── Follow-up email drafter ───────────────────────────────
                    _applied_at = job.get("applied_at") or ""
                    _days_since = None
                    if _applied_at:
                        try:
                            _applied_dt = datetime.fromisoformat(_applied_at[:19])
                            _days_since = (datetime.now() - _applied_dt).days
                        except Exception:
                            pass

                    if job["status"] == "applied" and _days_since is not None and _days_since >= 7:
                        st.divider()
                        st.caption(f"Applied {_days_since} days ago — no response yet.")
                        followup_key = f"followup_{job['id']}"
                        if st.button("✉️ Draft Follow-up Email", key=f"followup_btn_{job['id']}"):
                            _profile = load_profile()
                            with st.spinner("Drafting follow-up email…"):
                                try:
                                    followup_text = claude_call(
                                        system=(
                                            "You are helping a job candidate write a brief, professional follow-up email "
                                            "after submitting a job application with no response. "
                                            "The email should be:\n"
                                            "- 3-4 sentences max. Short. Respectful of their time.\n"
                                            "- Reference the specific role and company by name.\n"
                                            "- Reaffirm genuine interest without sounding desperate.\n"
                                            "- End with a clear, low-pressure call to action.\n"
                                            "- Sound like a real human wrote it — no buzzwords, no corporate filler.\n"
                                            "- No subject line — just the email body.\n"
                                            "- Sign off with the candidate's name."
                                        ),
                                        user=(
                                            f"Write a follow-up email for:\n"
                                            f"Role: {job['title']}\n"
                                            f"Company: {job['company']}\n"
                                            f"Applied: {_days_since} days ago\n"
                                            f"Candidate name: {_profile.get('name', '')}\n"
                                            f"Candidate email: {_profile.get('email', '')}"
                                        ),
                                        model="claude-haiku-4-5-20251001",
                                        max_tokens=300,
                                    )
                                    st.session_state[followup_key] = followup_text
                                except Exception as e:
                                    st.error(f"Follow-up draft failed: {e}")

                        if followup_key in st.session_state:
                            st.markdown("**✉️ Follow-up Draft**")
                            st.text_area(
                                "Copy and send:",
                                value=st.session_state[followup_key],
                                height=160,
                                key=f"followup_text_{job['id']}",
                            )
                            if st.button("Clear", key=f"followup_clear_{job['id']}"):
                                del st.session_state[followup_key]
                                st.rerun()

                    # ── Interview Prep ────────────────────────────────────────
                    st.divider()
                    prep_key = f"interview_prep_{job['id']}"
                    if prep_key not in st.session_state and job.get("interview_prep"):
                        st.session_state[prep_key] = job["interview_prep"]
                    prep_desc = job.get("description", "").strip()
                    if not prep_desc:
                        st.caption("🎤 Interview prep requires a job description — open the job on Job Board and paste one first.")
                    else:
                        if st.button("🎤 Generate Interview Prep", key=f"prep_btn_{job['id']}"):
                            _profile = load_profile()
                            _resume  = _profile.get("resume", "").strip()
                            with st.spinner("Claude is generating interview prep…"):
                                try:
                                    prep_raw = claude_call(
                                        system=(
                                            "You are an expert technical interview coach.\n"
                                            "Given a job description and the candidate's resume, generate the 10 most likely "
                                            "interview questions for this specific role.\n\n"
                                            "Return ONLY valid JSON — no markdown fences, no extra text:\n"
                                            '{"questions": [\n'
                                            '  {\n'
                                            '    "q": "<question text>",\n'
                                            '    "difficulty": "<easy|medium|hard>",\n'
                                            '    "answer": "<tailored answer using candidate\'s specific projects, metrics, and technologies — 3-5 sentences>",\n'
                                            '    "follow_ups": ["<follow-up question 1>", "<follow-up question 2>"]\n'
                                            '  }\n'
                                            "]}\n\n"
                                            "Rules:\n"
                                            "- Questions must be specific to THIS role and company, not generic.\n"
                                            "- Answers must reference the candidate's actual experience — real project names, real numbers.\n"
                                            "- Mix technical (4), behavioural (3), and situational (3) questions.\n"
                                            "- difficulty: easy=general/intro, medium=core technical, hard=deep-dive/system-design.\n"
                                            "- 2 follow-up questions per question — what an interviewer would probe next.\n"
                                            "- Do not invent experience the candidate doesn't have."
                                        ),
                                        user=(
                                            f"## Job Posting\n"
                                            f"Title: {job['title']}\n"
                                            f"Company: {job['company']}\n\n"
                                            f"{prep_desc[:4000]}\n\n"
                                            f"## Candidate Resume\n{_resume[:3000]}"
                                        ),
                                        model="claude-sonnet-4-6",
                                        max_tokens=4000,
                                    )
                                    # Validate JSON — fall back to storing raw if parse fails
                                    try:
                                        import json as _json
                                        _json.loads(prep_raw)
                                    except Exception:
                                        pass  # store as-is; display will fall back to markdown
                                    st.session_state[prep_key] = prep_raw
                                    update_interview_prep(job["id"], prep_raw)
                                except Exception as e:
                                    st.error(f"Interview prep failed: {e}")

                        if prep_key in st.session_state:
                            _prep_raw = st.session_state[prep_key]
                            _fname_base = f"interview_prep_{job['company'].replace(' ','_')}_{job['title'].replace(' ','_')[:30]}"
                            _doc = _Document()
                            _doc.add_heading(f"Interview Prep — {job['title']} @ {job['company']}", level=1)

                            # Try to parse as structured JSON (v2 format)
                            _questions = None
                            try:
                                import json as _json
                                _parsed = _json.loads(_prep_raw)
                                _questions = _parsed.get("questions", [])
                            except Exception:
                                pass

                            _DIFF_COLOR = {"easy": "#22c55e", "medium": "#f59e0b", "hard": "#ef4444"}

                            with st.expander("🎤 Interview Prep", expanded=True):
                                if _questions:
                                    for i, item in enumerate(_questions, 1):
                                        diff  = item.get("difficulty", "medium")
                                        color = _DIFF_COLOR.get(diff, "#94a3b8")
                                        st.markdown(
                                            f'<div style="margin-top:14px;">'
                                            f'<span style="font-weight:600;font-size:1rem;color:#f1f5f9;">Q{i}: {item["q"]}</span>'
                                            f'&nbsp;<span style="background:{color}22;color:{color};padding:1px 8px;'
                                            f'border-radius:99px;font-size:0.75rem;">{diff}</span>'
                                            f'</div>',
                                            unsafe_allow_html=True,
                                        )
                                        st.markdown(item.get("answer", ""))
                                        fups = item.get("follow_ups", [])
                                        if fups:
                                            st.caption("Follow-up questions the interviewer might ask:")
                                            for fup in fups:
                                                st.markdown(f"&nbsp;&nbsp;↳ *{fup}*")

                                        # Add to docx
                                        p = _doc.add_paragraph()
                                        run = p.add_run(f"Q{i} [{diff.upper()}]: {item['q']}")
                                        run.bold = True
                                        run.font.size = _Pt(11)
                                        _doc.add_paragraph(item.get("answer", ""))
                                        if fups:
                                            _doc.add_paragraph("Follow-ups:")
                                            for fup in fups:
                                                _doc.add_paragraph(f"  ↳ {fup}")
                                        _doc.add_paragraph("")
                                else:
                                    # Fallback: render as plain markdown (old format)
                                    st.markdown(_prep_raw)
                                    for line in _prep_raw.splitlines():
                                        line = line.strip()
                                        if not line:
                                            _doc.add_paragraph("")
                                        elif line.startswith("**Q") and line.endswith("**"):
                                            p = _doc.add_paragraph()
                                            run = p.add_run(line.strip("*"))
                                            run.bold = True
                                            run.font.size = _Pt(11)
                                        else:
                                            _doc.add_paragraph(line.lstrip("*").rstrip("*"))

                                dl_col, clr_col = st.columns([2, 1])
                                _buf = io.BytesIO()
                                _doc.save(_buf)
                                _buf.seek(0)
                                dl_col.download_button(
                                    "⬇ Download as .docx",
                                    data=_buf.getvalue(),
                                    file_name=f"{_fname_base}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"prep_dl_{job['id']}",
                                )
                                if has_libreoffice():
                                    import tempfile, os as _os
                                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as _tmp:
                                        _tmp_path = _tmp.name
                                        _doc.save(_tmp_path)
                                    _pdf_path = docx_to_pdf(_tmp_path)
                                    if _pdf_path and _os.path.exists(_pdf_path):
                                        with open(_pdf_path, "rb") as _pf:
                                            dl_col.download_button(
                                                "⬇ Download as .pdf",
                                                data=_pf.read(),
                                                file_name=f"{_fname_base}.pdf",
                                                mime="application/pdf",
                                                key=f"prep_dl_pdf_{job['id']}",
                                            )
                                if clr_col.button("Clear", key=f"prep_clear_{job['id']}"):
                                    del st.session_state[prep_key]
                                    update_interview_prep(job["id"], "")
                                    st.rerun()
